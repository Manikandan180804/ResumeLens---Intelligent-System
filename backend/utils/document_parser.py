"""
Document parsing utilities for PDF, DOCX, and TXT files.
"""
import io
import re
import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF file bytes."""
    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"PyPDF2 failed: {e}")
        try:
            from pdfminer.high_level import extract_text as pdfminer_extract
            return pdfminer_extract(io.BytesIO(file_bytes))
        except Exception as e2:
            logger.error(f"pdfminer also failed: {e2}")
            return ""


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX file bytes."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        # Also extract table text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except Exception as e:
        logger.error(f"DOCX parsing failed: {e}")
        return ""


def extract_text_from_image(file_bytes: bytes) -> str:
    """Extract text from Image file bytes using EasyOCR."""
    try:
        import easyocr
        import numpy as np
        from PIL import Image
        
        # Load image with PIL
        image = Image.open(io.BytesIO(file_bytes))
        # Convert to numpy array
        image_np = np.array(image)
        
        # Initialize reader (downloads models if missing)
        reader = easyocr.Reader(['en'], gpu=False) # Default to CPU to be safe
        results = reader.readtext(image_np)
        
        return "\n".join([res[1] for res in results])
    except Exception as e:
        logger.error(f"Image OCR failed: {e}")
        return ""


def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """Extract text based on file extension."""
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext in [".docx", ".doc"]:
        return extract_text_from_docx(file_bytes)
    elif ext in [".png", ".jpg", ".jpeg", ".bmp", ".tiff"]:
        return extract_text_from_image(file_bytes)
    elif ext == ".txt":
        return file_bytes.decode("utf-8", errors="ignore")
    else:
        # Try to decode as text
        return file_bytes.decode("utf-8", errors="ignore")


def clean_text(text: str) -> str:
    """Clean and normalize extracted text."""
    if not text:
        return ""
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)
    # Remove special chars but keep useful punctuation
    text = re.sub(r'[^\w\s\.\,\;\:\!\?\@\#\&\+\-\/\(\)\[\]\'\"]', ' ', text)
    return text.strip()
